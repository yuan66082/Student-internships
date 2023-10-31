# 这里面的东西不懂Python不用改，先看【说明.md】

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Mm,Pt,RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL

from 设置 import *

# 下面这两个导入是使用函数设置单元格框线需要的类
from docx.table import _Cell
from docx.oxml import OxmlElement

# 以下函数为生成docx文档的函数
def make_docx(obj):
    doc = Document()
    # 设置页面宽高为A4
    doc.sections[0].page_width = Mm(210)
    doc.sections[0].page_height = Mm(297)
    # 设置默认样式的字体、字号、颜色、段后间距
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Normal'].font.size = Pt(12)
    doc.styles['Normal'].font.color.rgb = RGBColor(0,0,0)
    doc.styles['Normal'].paragraph_format.space_after = 0

    # 添加两个‘字体’样式，一个是标题的样式，一个是加粗
    doc.styles.add_style('Header1', WD_STYLE_TYPE.CHARACTER)
    doc.styles.add_style('Bold1', WD_STYLE_TYPE.CHARACTER)
    # 设置加粗样式
    doc.styles['Bold1'].font.name = u'宋体'
    doc.styles['Bold1']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Bold1'].font.size = Pt(12)
    doc.styles['Bold1'].font.color.rgb = RGBColor(0,0,0)
    doc.styles['Bold1'].font.bold = True
    # 设置标题样式
    doc.styles['Header1'].font.name = u'宋体'
    doc.styles['Header1']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Header1'].font.size = Pt(20)
    doc.styles['Header1'].font.color.rgb = RGBColor(0,0,0)


    # 下面这段函数是设置单元格框线，没办法，python-docx没有这个功能，网上下的，直接用--------------------------------------------------------
    # 定义一个默认的边框a 1px 单框线 黑色
    a = {"sz": 8, "val": "single", "color": "#000000", "space": "0"}
    def set_cell_border(cell: _Cell, **kwargs):
        """
        设置单元格边框
        用法:
        set_cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
            bottom={"sz": 12, "color": "#00FF00", "val": "single"},
            start={"sz": 24, "val": "dashed", "shadow": "true"},
            end={"sz": 12, "val": "dashed"},
        )
        """

        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
    
        # check for tag existnace, if none found, then create one
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
    
        # list over all available tags
        for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)
    
                # check for tag existnace, if none found, then create one
                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)
    
                # looks like order of attributes is important
                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))
    # ------------------------------------------------------------------------------------

    # 开始写文档，直接添加一个一行四列的表格
    table = doc.add_table(rows=1, cols=4)

    # 第一行，先合并单元格
    table.cell(0,0).merge(table.cell(0,3))
    # 设置行高
    table.rows[0].height = Mm(10.2)
    # 设置单元格垂直居中
    table.cell(0,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 添加单元格的内容，在paragraphs[0]中加入两个run，因为字体样式不一样
    table.rows[0].cells[0].paragraphs[0].add_run('附件2：       ', doc.styles['Bold1'])
    table.rows[0].cells[0].paragraphs[0].add_run(学生届别 + '学生实习指导记录表', doc.styles['Header1'])

    # 第二行
    table.add_row()
    # 设置行高
    table.rows[1].height = Mm(10.2)
    # 单元格内垂直居中
    table.cell(1,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 合并单元格
    table.cell(1,0).merge(table.cell(1,3))
    # 添加单元格的内容，设置为加粗样式
    table.rows[1].cells[0].paragraphs[0].add_run('  二级学院：' + 二级学院, doc.styles['Bold1'])

    # 第三行
    # 添加行
    table.add_row()
    # 设置行高
    table.rows[2].height = Mm(10.2)
    # 在第一格和第三格填入内容
    table.rows[2].cells[0].paragraphs[0].add_run('指导教师')
    table.rows[2].cells[1].paragraphs[0].add_run(指导教师)
    table.rows[2].cells[2].paragraphs[0].add_run('指导学生')
    table.rows[2].cells[3].paragraphs[0].add_run(obj['姓名'])
    # 循环所有格
    for cell in table.rows[2].cells:
        # 四面边框
        set_cell_border(cell,top=a,bottom=a,start=a,end=a)
        # 单元格内垂直居中
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 单元格内的段落水平居中
        cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

    # 第四行
    # 添加行
    table.add_row()
    # 设置行高
    table.rows[3].height = Mm(10.2)
    # 在第一格和第三格填入内容
    table.rows[3].cells[0].paragraphs[0].add_run('专业班级')
    table.rows[3].cells[1].paragraphs[0].add_run(专业班级)
    table.rows[3].cells[2].paragraphs[0].add_run('学生辅导员')
    table.rows[3].cells[3].paragraphs[0].add_run(学生辅导员)
    # 循环所有格
    for cell in table.rows[3].cells:
        # 四面边框
        set_cell_border(cell,top=a,bottom=a,start=a,end=a)
        # 单元格内垂直居中
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 单元格内的段落水平居中
        cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

    # 第五行
    # 添加行
    table.add_row()
    # 设置行高
    table.rows[4].height = Mm(10.2)
    # 在第一格和第三格填入内容
    table.rows[4].cells[0].paragraphs[0].add_run('实习企业')
    table.rows[4].cells[1].paragraphs[0].add_run(obj['公司'])
    table.rows[4].cells[2].paragraphs[0].add_run('岗位')
    table.rows[4].cells[3].paragraphs[0].add_run(obj['岗位'])
    # 循环所有格
    for cell in table.rows[4].cells:
        # 四面边框
        set_cell_border(cell,top=a,bottom=a,start=a,end=a)
        # 单元格内垂直居中
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 单元格内的段落水平居中
        cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

    # 第六行
    # 添加行
    table.add_row()
    # 设置行高
    table.rows[5].height = Mm(10.2)
    # 合并2，3，4格
    table.cell(5,1).merge(table.cell(5,3))
    # 在第一格和第三格填入内容
    table.rows[5].cells[0].paragraphs[0].add_run('实习地点')
    table.rows[5].cells[1].paragraphs[0].add_run(obj['公司'])
    # 循环所有格
    for cell in table.rows[5].cells:
        # 四面边框
        set_cell_border(cell,top=a,bottom=a,start=a,end=a)
        # 单元格内垂直居中
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 单元格内的段落水平居中
        cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

    # 第七行
    table.add_row()
    # 设置行高
    table.rows[6].height = Mm(11.8)
    # 合并1，2，3，4格
    table.cell(6,0).merge(table.cell(6,3))
    # 单元格内垂直居中
    table.cell(6,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 填写内容
    table.rows[6].cells[0].paragraphs[0].add_run('实习指导记录', doc.styles['Bold1'])
    # 水平居中
    table.rows[6].cells[0].paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER
    # 四面框线
    set_cell_border(table.rows[6].cells[0],top=a,bottom=a,start=a,end=a)

    # 第八行
    # 添加行
    table.add_row()
    # 设置行高
    table.rows[7].height = Mm(10.2)
    # 合并2，3，4格
    table.cell(7,1).merge(table.cell(7,3))
    # 在第一格和第三格填入内容
    table.rows[7].cells[0].paragraphs[0].add_run('指导时间')
    table.rows[7].cells[1].paragraphs[0].add_run('指导内容')
    # 循环所有格
    for cell in table.rows[7].cells:
        # 四面边框
        set_cell_border(cell,top=a,bottom=a,start=a,end=a)
        # 单元格内垂直居中
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 单元格内的段落水平居中
        cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER
    
    # 第九行开始写指导记录，根据指导记录的条数添加行
    for i in range(8,len(指导记录)+8):
        # 添加行
        table.add_row()
        # 设置行高
        table.rows[i].height = Mm(30)
        # 合并2，3，4格
        table.cell(i,1).merge(table.cell(i,3))
        # 在第一格和第二格填入内容
        table.rows[i].cells[0].paragraphs[0].add_run(指导记录[i-8]['时间'])
        table.rows[i].cells[1].paragraphs[0].add_run(指导记录[i-8]['内容'])
        # 循环所有格
        for cell in table.rows[i].cells:
            # 四面边框
            set_cell_border(cell,top=a,bottom=a,start=a,end=a)
            # 单元格内垂直居中
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # 单元格内的段落水平居中
            cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

    '''
    # 第九行 开始填写指导记录
    # 添加行
    table.add_row()
    # 设置行高
    table.rows[8].height = Mm(30)
    # 合并2，3，4格
    table.cell(8,1).merge(table.cell(8,3))
    # 在第一格和第三格填入内容
    table.rows[8].cells[0].paragraphs[0].add_run(指导记录[0]['时间'])
    table.rows[8].cells[1].paragraphs[0].add_run(指导记录[0]['内容'])
    # 循环所有格
    for cell in table.rows[8].cells:
        # 四面边框
        set_cell_border(cell,top=a,bottom=a,start=a,end=a)
        # 单元格内垂直居中
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 单元格内的段落水平居中
        cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

    # 第十行 第二条指导记录
    # 添加行
    table.add_row()
    # 设置行高
    table.rows[9].height = Mm(30)
    # 合并2，3，4格
    table.cell(9,1).merge(table.cell(9,3))
    # 在第一格和第三格填入内容
    table.rows[9].cells[0].paragraphs[0].add_run(指导记录[1]['时间'])
    table.rows[9].cells[1].paragraphs[0].add_run(指导记录[1]['内容'])
    # 循环所有格
    for cell in table.rows[9].cells:
        # 四面边框
        set_cell_border(cell,top=a,bottom=a,start=a,end=a)
        # 单元格内垂直居中
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 单元格内的段落水平居中
        cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

    # 第十一行 第三条指导记录
    # 添加行
    table.add_row()
    # 设置行高
    table.rows[10].height = Mm(30)
    # 合并2，3，4格
    table.cell(10,1).merge(table.cell(10,3))
    # 在第一格和第三格填入内容
    table.rows[10].cells[0].paragraphs[0].add_run(指导记录[2]['时间'])
    table.rows[10].cells[1].paragraphs[0].add_run(指导记录[2]['内容'])
    # 循环所有格
    for cell in table.rows[10].cells:
        # 四面边框
        set_cell_border(cell,top=a,bottom=a,start=a,end=a)
        # 单元格内垂直居中
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 单元格内的段落水平居中
        cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

    # 第十二行 第四条指导记录
    # 添加行
    table.add_row()
    # 设置行高
    table.rows[11].height = Mm(30)
    # 合并2，3，4格
    table.cell(11,1).merge(table.cell(11,3))
    # 在第一格和第三格填入内容
    table.rows[11].cells[0].paragraphs[0].add_run(指导记录[3]['时间'])
    table.rows[11].cells[1].paragraphs[0].add_run(指导记录[3]['内容'])
    # 循环所有格
    for cell in table.rows[11].cells:
        # 四面边框
        set_cell_border(cell,top=a,bottom=a,start=a,end=a)
        # 单元格内垂直居中
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 单元格内的段落水平居中
        cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER
    '''
    
    # 空行
    # 判断用户输入的指导记录是不是超过了12条，没超过就补空行到第20行，超过了就不补
    if len(指导记录) < 12:
        for i in range(8+len(指导记录),20):
            # 添加行
            table.add_row()
            # 设置行高
            table.rows[i].height = Mm(30)
            # 合并2，3，4格
            table.cell(i,1).merge(table.cell(i,3))
            # 在第一格和第三格填入内容
            # table.rows[i].cells[0].paragraphs[0].add_run(指导记录[3]['时间'])
            # table.rows[i].cells[1].paragraphs[0].add_run(指导记录[3]['内容'])
            # 循环所有格
            for cell in table.rows[i].cells:
                # 四面边框
                set_cell_border(cell,top=a,bottom=a,start=a,end=a)
                # 单元格内垂直居中
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                # 单元格内的段落水平居中
                cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER


    # 最后一行
    table.add_row()
    table.cell(len(table.rows)-1,0).merge(table.cell(len(table.rows)-1,3))
    # 设置行高
    table.rows[len(table.rows)-1].height = Mm(10.2)
    table.cell(len(table.rows)-1,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.rows[len(table.rows)-1].cells[0].paragraphs[0].add_run('备注：原则上一个学生一张表，每月最少记录两次，分行逐一填写。')

    doc.save(r'./layout/' + 学生届别 + '学生实习指导记录表_'+ obj["姓名"] + '.docx')





